import streamlit as st
from docx import Document
from datetime import date

def create_report(data):
    doc = Document()
    doc.add_heading('【施設仕様版BPSDケアレポート】', level=1)

    doc.add_heading('Ⅰ. 施設情報', level=2)
    doc.add_paragraph(f"施設名：{data['facility_name']}\n部署名：{data['department']}\n担当者：{data['staff_name']}\n作成日：{data['created_date']}")

    doc.add_heading('Ⅱ. 基本情報', level=2)
    doc.add_paragraph(f"名前：{data['name']}\n年齢：{data['age']}\n性別：{data['gender']}\n生活歴・背景：{data['background']}\n支援者：{data['supporter']}")

    doc.add_heading('Ⅲ. 7項目チェックリスト結果', level=2)
    for i, answer in enumerate(data['checklist'], 1):
        doc.add_paragraph(f"{i}　{answer}")

    doc.add_heading('Ⅳ. BPSD25Q結果まとめ', level=2)
    for item, score in data['bpsd'].items():
        doc.add_paragraph(f"{item}：{score}点")

    doc.add_heading('Ⅴ. アセスメントと仮説', level=2)
    doc.add_paragraph(data['assessment'])

    doc.add_heading('Ⅵ. ケア方針・支援計画', level=2)
    doc.add_paragraph(data['care_plan'])

    doc.add_heading('Ⅶ. モニタリング・評価計画', level=2)
    doc.add_paragraph(data['monitoring'])

    doc.add_heading('Ⅷ. 署名欄', level=2)
    doc.add_paragraph(f"実施責任者：{data['responsible']}\nケアマネジャ：{data['care_manager']}\nチームリーダー：{data['team_leader']}")

    return doc

st.title('【認知症BPSD支援用レポート自動作成フォーム】')

st.header('1. 施設情報')
facility_name = st.text_input('施設名')
department = st.text_input('部署名')
staff_name = st.text_input('担当者名')
created_date = st.date_input('作成日', value=date.today())

st.header('2. 基本情報')
name = st.text_input('対象者氏名')
age = st.number_input('年齢', min_value=0, max_value=120, step=1)
gender = st.selectbox('性別', ['男性', '女性'])
background = st.text_area('生活歴・背景')
supporter = st.text_input('主な支援者')

st.header('3. 7項目チェックリスト')
checklist = [st.text_input(f'項目{i}') for i in range(1, 8)]

st.header('4. BPSD25Q簡易評価')
bpsd_items = ['幻視・幻聴', '妄想', '暴言', '暴行', '徘徊・不穏', '無断外出', '性的不適切行動',
              '常同行動', '脱抑制', '怒りっぽい', '繰り返し質問', '収集', '大声',
              '抑うつ', '無気力', '無反応・無関心', '不安', '傾眠傾向', '閉じこもり',
              '昼夜逆転', '異食・過食', '介護への抵抗', '不潔行為', '火の不始末', '物をなくす']
bpsd = {item: st.selectbox(item, options=list(range(6)), index=0) for item in bpsd_items}

st.header('5. アセスメント仮説')
assessment = st.text_area('背景要因・潜在ニーズ等')

st.header('6. ケア方針・支援計画')
care_plan = st.text_area('具体的なケア内容')

st.header('7. モニタリング・評価計画')
monitoring = st.text_area('モニタリングと評価方法')

st.header('8. 署名欄')
responsible = st.text_input('実施責任者名')
care_manager = st.text_input('ケアマネジャー名')
team_leader = st.text_input('チームリーダー名')

if st.button('レポート作成'):
    data = {
        'facility_name': facility_name,
        'department': department,
        'staff_name': staff_name,
        'created_date': created_date,
        'name': name,
        'age': age,
        'gender': gender,
        'background': background,
        'supporter': supporter,
        'checklist': checklist,
        'bpsd': bpsd,
        'assessment': assessment,
        'care_plan': care_plan,
        'monitoring': monitoring,
        'responsible': responsible,
        'care_manager': care_manager,
        'team_leader': team_leader
    }
    doc = create_report(data)
    file_path = '/mnt/data/bpsd_care_report.docx'
    doc.save(file_path)
    st.success('レポートが作成されました！')
    with open(file_path, 'rb') as f:
        st.download_button('レポートをダウンロードする', f, file_name='BPSD_Care_Report.docx')
